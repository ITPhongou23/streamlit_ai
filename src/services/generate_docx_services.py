from datetime import datetime
from io import BytesIO

from docx import Document


def generate_docx(content, result_label, score):
    doc = Document()
    doc.add_heading('AI News Indicator - Báo cáo kết quả', 0)
    doc.add_paragraph(
        f"Thời gian tạo: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
    )
    doc.add_heading('Kết quả phân tích', level=1)
    doc.add_paragraph(f"Kết quả: {result_label}")
    doc.add_paragraph(f"Mức độ tin cậy: {score * 100:.2f}%")
    doc.add_heading('Nội dung văn bản input', level=1)
    doc.add_paragraph(content)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
